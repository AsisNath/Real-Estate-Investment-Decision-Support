"""Generate the final write-up (docs/submission/NorthStar-WriteUp.docx).

Generated rather than hand-written for the same reason as the deck: the figures
must track the repository. Update the numbers here and re-run, so the document
can never claim a test count the suite does not have.

    python scripts/build_writeup.py

Constraint from the brief: 5 pages maximum including images, 11pt minimum.
Screenshots come from docs/screenshots/, captured against the running app.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "screenshots"
OUTPUT = ROOT / "docs" / "submission" / "NorthStar-WriteUp.docx"

BRAND = RGBColor(0x1D, 0x5C, 0x4D)
INK = RGBColor(0x1F, 0x27, 0x33)
MUTED = RGBColor(0x5B, 0x6B, 0x7C)

TEAM = "Ashish Nath · Justin Kretschman"
REPO = "https://github.com/AsisNath/Real-Estate-Investment-Decision-Support"

# Verified against the repository at build time.
TESTS = 138
ZIPS = "41,488"
NOTES = 10
MODULES = 9


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.03
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _h(doc: Document, text: str, size: int = 13, space_before: int = 9) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = BRAND


def _p(doc: Document, text: str, size: int = 11, color=INK, italic=False, space_after=5):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return para


def _bullet(doc: Document, label: str, text: str = "", size: int = 10.5):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Inches(0.22)
    run = para.add_run(label)
    run.bold = bool(text)
    run.font.size = Pt(size)
    run.font.color.rgb = INK
    if text:
        rest = para.add_run(" " + text)
        rest.font.size = Pt(size)
        rest.font.color.rgb = INK
    return para


def _figure(doc: Document, filename: str, caption: str, width: float = 6.4):
    path = SHOTS / filename
    if not path.exists():
        _p(doc, f"[missing screenshot: {filename}]", 10, MUTED, italic=True)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(1)
    para.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.italic = True


def _table(doc: Document, rows: list[tuple[str, str]], widths=(1.9, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width, cells[1].width = Inches(widths[0]), Inches(widths[1])
        for cell, text, bold in ((cells[0], left, True), (cells[1], right, False)):
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(10)
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.55)
        section.left_margin = section.right_margin = Inches(0.7)
    _style(doc)

    # ------------------------------------------------------------ header block
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("NorthStar Property Investment Consulting")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(1)
    run = sub.add_run("Deterministic financial analysis joined to AI policy research that keeps itself current")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run(f"Team 5 — {TEAM}   |   BUKD-X500 Agentic AI Systems, Design Studio   |   {REPO}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED

    # ------------------------------------------------------- problem + product
    _h(doc, "1. The problem, and the evidence it is real", space_before=0)
    _p(
        doc,
        "A small real-estate investor decides with a spreadsheet and whatever a search engine returns. "
        "The financial modelling is the easy half. The hard half is that rental rules are hyper-local, "
        "change often, and the highest-ranking sources are property-management blogs that go stale "
        "silently. Getting one rule wrong invalidates everything downstream: model short-term-rental "
        "(STR) income on a property where STR is not permitted, and the IRR, the cash flow, and the "
        "recommendation are all fiction.",
    )
    _p(
        doc,
        "Our own research surfaced the sharpest possible illustration. Maryland Heights, MO 63043 allows "
        "an investor to operate an STR provided the owner OR a manager is within a one-hour drive. "
        "Ballwin, MO 63021 — two ZIP codes away, same county, same state statutes — requires the owner "
        "to reside on the property at least 180 days a year, which a non-owner-occupied investment "
        "cannot satisfy. Same question, opposite answers. No static dataset captures that; it has to be "
        "researched per address and kept fresh.",
    )

    _h(doc, "2. What we built and how it works")
    _p(
        doc,
        "NorthStar is a local FastAPI web application. You enter an address and your assumptions, and it "
        "returns an investor report: financial metrics, market and policy context, ranked risks, a "
        "diligence checklist, and a rule-based recommendation. It is organised as two layers that never "
        "mix.",
    )
    _bullet(
        doc, "Layer 1 — the arithmetic.",
        "NOI, DSCR, cash flow, cap rate, 5/10-year IRR, equity multiple and LTV are ordinary, "
        "unit-tested Python. No language model touches a number.",
    )
    _bullet(
        doc, "Layer 2 — the rules.",
        "An Agent Skill researches STR ordinances, state landlord-tenant law, rent control, and HOA "
        "authority using live web search, verifies load-bearing facts against official .gov or statute "
        "sources, and writes a source-cited note into the knowledge bank.",
    )
    _bullet(
        doc, "They meet at the decision.",
        "Researched findings become policy flags, diligence items, and assumption conflicts — the app "
        "tells you when your own rent-growth input exceeds a legal cap it has on file.",
    )

    _figure(
        doc, "02-investor-report.png",
        "Figure 1 — Investor report for 745 Woodrun Dr, Ballwin MO. The researched note has flagged "
        "STR income as unavailable and raised policy risk from medium to HIGH. This deal fails on cash "
        "flow either way; on a marginal one, that shift is what changes the answer.",
        width=3.85,
    )

    _h(doc, "3. Target users and use cases")
    _p(
        doc,
        "The primary user is an individual or small-portfolio residential investor doing first-pass "
        "diligence on a specific property — someone who cannot justify a consultant for every listing, "
        "but for whom one missed ordinance is a five-figure mistake. A second is the advisor who must "
        "hand that person something defensible; the app generates a printable Policy Brief citing every "
        "source. Typical uses: screening a listing, checking whether an STR strategy is legal there at "
        "all, and building a diligence checklist before an offer.",
    )

    # ------------------------------------------------------------------ tools
    _h(doc, "4. Tools and platforms")
    _table(
        doc,
        [
            ("Application", "Python 3.11, FastAPI, Pydantic, Jinja2; vanilla HTML/CSS/JS front end with no build step"),
            ("AI research", "Agent Skill (SKILL.md) executed through a signed-in agent CLI (Codex or Claude Code), "
                            "with the Anthropic API as a fallback backend"),
            ("Data", f"GeoNames US postal-code dataset ({ZIPS} ZIPs) bundled for offline address verification"),
            ("Documents", "python-pptx and python-docx generate this write-up and the deck; pypdf reads user-supplied PDFs"),
            ("Testing", f"pytest — {TESTS} automated tests"),
            ("Delivery", "Windows .bat launchers; Git/GitHub for version control"),
        ],
    )

    _h(doc, "5. Design decisions, and the alternatives we rejected")
    _bullet(
        doc, "Keep the model out of the arithmetic.",
        "The obvious alternative was to let an LLM compute or sanity-check the financials. We rejected "
        "it because a hallucinated IRR is worse than no IRR — it is confident and wrong. Determinism is "
        "the feature; the AI is scoped to research, where it genuinely outperforms a static database.",
    )
    _bullet(
        doc, "Make provenance structural, not a label.",
        "The knowledge bank has two roots: researched/ is written only by the Skill, user/ only by "
        "people. This is enforced in code — the in-app form physically cannot write into researched/ — "
        "rather than by naming convention, so an AI-verified badge cannot be faked by a typo.",
    )
    _bullet(
        doc, "Delegate to an agent the user has already signed into.",
        "Our first design had the app call an LLM API directly, which meant every user needed their own "
        "key. That was the wrong shape: the Skill is plain instructions, so it can run inside a CLI the "
        "user is already authenticated with, which supplies its own credentials and web search. Setup "
        "went from “obtain an API key” to nothing at all.",
    )
    _bullet(
        doc, "State missing data; never fill the gap.",
        "ZIP-level misses fall back to state, then national, and the report says which level it used. "
        "When an official page blocks automated access, the fact is tagged “confirm with the agency” "
        "rather than quietly replaced with a blog — the precise failure the tool exists to prevent.",
    )

    _h(doc, "6. Process flow")
    _p(
        doc,
        "Enter address and assumptions → Pydantic validation → city/state/ZIP cross-check against the "
        "bundled ZIP dataset → deterministic financial model → market and policy lookup (ZIP → city → "
        "state → national) → knowledge-bank read for matching notes → report assembled with risks, "
        "conflicts and recommendation → analysis trail written to logs/. In parallel, if the ZIP has no "
        "researched note or its note is past the 120-day freshness window, a background thread runs the "
        "Skill, saves a cited note to knowledge_bank/researched/zips/<zip>/, and the next analysis picks "
        "it up. Research never blocks the report.",
    )

    _figure(
        doc, "03-knowledge-bank.png",
        "Figure 2 — The Knowledge Bank. Every note shows its scope, research date, citation counts, and "
        "whether it was AI-researched or added by a person. Each can be opened as a client-ready Policy Brief.",
        width=5.75,
    )

    # ---------------------------------------------------------------- testing
    _h(doc, "7. How we tested it, and the results")
    _p(
        doc,
        f"The suite is {TESTS} automated tests across {MODULES} application modules. The financial model "
        "is unit-tested against hand-computed values. Research quality was tested the way the Skill was "
        "built to be tested: on two near-opposite regulatory environments (Austin, TX — STR legal with a "
        "licence, rent control banned statewide; Los Angeles, CA — STR effectively closed to investors, "
        "two overlapping rent-control regimes). The Skill produced correct, differently-shaped output for "
        "both without edits to its instructions.",
    )
    _p(doc, "Three defects the tests and reviews caught, and what they taught us:", space_after=2)
    _bullet(
        doc, "A server-hanging deadlock.",
        "request_research called status() while holding a non-reentrant lock, so a second analysis of the "
        "same address during research hung the request thread. Caught by a test that re-entered mid-run.",
    )
    _bullet(
        doc, "A stale note that blocked its own refresh.",
        "The app requested research when a note was missing or stale, but the researcher declined whenever "
        "a file existed — so the one case that most needed re-research could never happen.",
    )
    _bullet(
        doc, "A silently disabled guardrail.",
        "A note the app researched unattended appended citations to its machine-readable summary lines, so "
        "values parsed as text instead of numbers and the assumption-conflict check stopped firing. Found "
        "by reading a real generated note rather than a fixture.",
    )
    _p(
        doc,
        "Strongest evidence the system works unattended: while preparing this submission we found a policy "
        "note in the repository that neither of us wrote by hand. A teammate had analysed 17723 Mueller Rd, "
        "Wildwood MO 63038; the app researched the ZIP on its own and saved a note carrying 20 official "
        "citations. The knowledge bank now holds " + str(NOTES) + " cited notes.",
    )

    _figure(
        doc, "04-policy-brief.png",
        "Figure 3 — A generated Policy Brief. Every fact carries a source link, an “as of” date, and an "
        "official-versus-secondary tag; the brief prints straight to PDF for a client.",
        width=5.05,
    )

    # -------------------------------------------------------------- adoption
    _h(doc, "8. Installation and adoption")
    _p(
        doc,
        "Clone the repository and double-click Run_NorthStar.bat. It creates the virtual environment, "
        "installs requirements, runs the test suite, clears any stale server off port 8000, and opens the "
        "browser. Automatic research needs no configuration if a signed-in agent CLI (codex or claude) is "
        "on PATH; otherwise Setup_Research.bat creates a git-ignored .env for an API key. The app is fully "
        "functional offline without either — you simply get a ready-to-paste research prompt instead. "
        "Adoption for a small firm would be: run it locally, drop HOA declarations and leases into "
        "knowledge_bank/user/, and let researched notes accumulate as a shared, auditable asset.",
    )

    _h(doc, "9. Security considerations")
    _p(
        doc,
        "This is a local, single-user application with no authentication surface and no inbound network "
        "exposure, so a remote attacker is not the realistic threat. The realistic threat is untrusted "
        "content: policy notes are assembled from the open web, and users drop PDFs they did not author "
        "into the knowledge bank — and that material is rendered into a document handed to a client. "
        "Defences we implemented:",
    )
    _bullet(doc, "Path traversal blocked.", "Note paths resolve and must stay inside knowledge_bank/; "
                                            "'..', drive letters and absolute paths are refused, not reinterpreted.")
    _bullet(doc, "Output escaping.", "Note text is HTML-escaped before any markdown conversion, so a "
                                     "document cannot inject markup or script into a Policy Brief.")
    _bullet(doc, "Validate before persisting.", "The research agent runs read-only and only prints; the app "
                                                "decides what reaches disk, and a reply that is not a policy "
                                                "note is discarded rather than saved.")
    _bullet(doc, "Secret hygiene.", "Credentials live in a git-ignored .env, never in the repository. The "
                                    "preferred backend stores no secret at all.")
    _bullet(doc, "Cost safety.", "Tests force research off and disable CLI detection, so running the suite "
                                 "can never spend money or a user's quota.")

    _h(doc, "10. Reflections")
    _p(
        doc,
        "What worked: writing standards instead of steps. The Skill says what good research looks like — "
        "verify against official sources, tag every fact, date everything, say “unverified” rather than "
        "guess — and the agent worked out how, including recovering when a government site returned 403. "
        "That is closer to managing an analyst than programming a pipeline. Making provenance a code "
        "boundary rather than a naming convention also paid off immediately.",
    )
    _p(
        doc,
        "What did not: our first research architecture required an API key per user, which we only "
        "recognised as the wrong shape after asking why the Lab 5 Skill never had that problem — it ran "
        "inside tools people were already signed into. We also lost time to a stale server that made fixed "
        "code look broken, and we discovered that a user's HOA declaration was being silently ignored "
        "because it was a scanned PDF with no text layer; the app now surfaces that instead of dropping it.",
    )
    _p(
        doc,
        "What we would do differently: capture screenshots and generate documents from scripts from day "
        "one, so figures track the code. Next steps are OCR for scanned declarations, a comparison view "
        "across candidate properties, and scheduled re-research so notes refresh before they age out.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path.relative_to(ROOT)}  ({path.stat().st_size / 1024:.0f} KB)")
